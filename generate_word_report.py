import re
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_md_paragraph(doc, text, style=None):
    p = doc.add_paragraph(style=style)
    # Basic markdown parsing: **bold**, `code`, *italic*
    parts = re.split(r'(\*\*.*?\*\*|`.*?`|\*.*?\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = p.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = p.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
        else:
            p.add_run(part)

doc = Document()

# Set Page Size to A4 & Margins to Moderate
section = doc.sections[0]
section.page_width = Inches(8.27)
section.page_height = Inches(11.69)
section.top_margin = Inches(1.0)
section.bottom_margin = Inches(1.0)
section.left_margin = Inches(0.75)
section.right_margin = Inches(0.75)

# Cover Page (Title)
title = doc.add_heading('Heart Disease Prediction Project Report', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# Cover Page (Team Info)
doc.add_heading('Team Details', level=1)
p_team = doc.add_paragraph()
p_team.add_run('Team Name: ').bold = True
p_team.add_run('NULL\n')

# Team Member Table
members = [
    ("Youssef Wael Mohamed", "ID_Num"),
    ("Youssef Abdelmonem Abdallah Ahmed", "ID_Num"),
    ("Abdelrahman Ahmed Farouk Mohamed Fouad", "ID_Num"),
    ("Youssef Mohamed Sobhy Ghareeb", "ID_Num"),
    ("Abdelrahman Yasser Ahmed Mohamed", "ID_Num"),
    ("Mohanad Hossam Arafa Abdelaziz", "ID_Num"),
    ("Youssef Mahmoud Fawzy Abbas", "ID_Num")
]

table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Student Name'
hdr_cells[1].text = 'Academic ID'
# Make Header Bold
for cell in hdr_cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

for name, st_id in members:
    row_cells = table.add_row().cells
    row_cells[0].text = name
    row_cells[1].text = st_id

doc.add_page_break()

# Read the actual report
with open('Project_Report.md', 'r', encoding='utf-8') as f:
    lines = f.readlines()

in_table = False
table_obj = None

image_buffer = []

def flush_images():
    global image_buffer
    if not image_buffer:
        return
    
    # Group chest pain & decision boundary
    is_grouped = False
    if len(image_buffer) == 2:
        names = [os.path.basename(p) for p in image_buffer]
        if any('Chest_Pain' in n for n in names) and any('Decision_Boundary' in n for n in names):
            is_grouped = True
            
    if is_grouped:
        img_table = doc.add_table(rows=1, cols=2)
        img_table.autofit = False
        for idx, img_path in enumerate(image_buffer):
            cell = img_table.rows[0].cells[idx]
            p = cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if os.path.exists(img_path):
                p.add_run().add_picture(img_path, width=Inches(3.1))
            else:
                p.add_run().add_text(f"[Image Missing]")
    else:
        for img_path in image_buffer:
            if os.path.exists(img_path):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(img_path, width=Inches(6.0))
            else:
                doc.add_paragraph(f"[Image Missing: {img_path}]")
                
    image_buffer.clear()

for line in lines:
    line = line.strip()
    if not line:
        flush_images()
        in_table = False
        doc.add_paragraph()
        continue
        
    if line.startswith('# Heart Disease Prediction'):
        continue  # Skip main title since we already added it.
        
    if line.startswith('## '):
        flush_images()
        in_table = False
        doc.add_heading(line[3:], level=1)
    elif line.startswith('### '):
        flush_images()
        in_table = False
        doc.add_heading(line[4:], level=2)
    elif line.startswith('#### '):
        flush_images()
        in_table = False
        doc.add_heading(line[5:], level=3)
    elif line.startswith('!['):
        m = re.search(r'!\[.*?\]\((.*?)\)', line)
        if m:
            image_buffer.append(m.group(1))
    elif line.startswith('- '):
        flush_images()
        in_table = False
        add_md_paragraph(doc, line[2:], style='List Bullet')
    elif line.startswith('|'):
        flush_images()
        # Table parsing logic
        parts = [p.strip() for p in line.split('|') if p.strip()]
        if not parts: continue
        if all(re.match(r'^[-:]+$', p) for p in parts):
            continue # Separator line
            
        if not in_table:
            in_table = True
            table_obj = doc.add_table(rows=1, cols=len(parts))
            table_obj.style = 'Table Grid'
            hdr_cells = table_obj.rows[0].cells
            for i, p in enumerate(parts):
                if i < len(hdr_cells):
                    hdr_cells[i].text = p
                    for paragraph in hdr_cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
        else:
            row_cells = table_obj.add_row().cells
            for i, p in enumerate(parts):
                if i < len(row_cells):
                    row_cells[i].text = p
    else:
        flush_images()
        in_table = False
        if line == '---':
            doc.add_page_break()
        else:
            add_md_paragraph(doc, line)

flush_images()

doc.save('Team_NULL_Project_Report.docx')
print("Successfully generated Word Document: Team_NULL_Project_Report.docx")
